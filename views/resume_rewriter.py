"""
Resume Rewriter View
"""

import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from src.integrations.resume_rewriter import ResumeRewriterIntegration


# ==========================================
# Create Word Document
# ==========================================

def create_resume_docx(resume_text):

    document = Document()

    # Set default font
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = resume_text.split("\n")

    first_line = True

    for line in lines:

        line = line.strip()

        if not line:
            continue

        # ----------------------------------
        # Candidate Name (first line)
        # ----------------------------------

        if first_line:

            heading = document.add_heading(level=0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run = heading.add_run(line)
            run.bold = True

            first_line = False
            continue

        # ----------------------------------
        # Detect Section Heading
        # ----------------------------------

        if (
            line.isupper()
            or line.endswith(":")
            or line.lower()
            in [
                "professional summary",
                "technical skills",
                "skills",
                "projects",
                "experience",
                "work experience",
                "education",
                "certifications",
                "achievements",
                "internships",
                "extracurricular activities",
            ]
        ):

            document.add_heading(line.replace(":", ""), level=2)

        # ----------------------------------
        # Bullet Point
        # ----------------------------------

        elif line.startswith("•") or line.startswith("-"):

            document.add_paragraph(
                line.replace("•", "").replace("-", "").strip(),
                style="List Bullet",
            )

        # ----------------------------------
        # Normal Paragraph
        # ----------------------------------

        else:

            document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer


# ==========================================
# View
# ==========================================

def show_resume_rewriter():

    st.title("✍️ Resume Rewriter")

    st.markdown(
        """
Rewrite your resume into a clean, professional, ATS-friendly resume while preserving all factual information.
"""
    )

    st.divider()

    # ----------------------------------
    # Validation
    # ----------------------------------

    if not st.session_state.get("resume_uploaded", False):

        st.warning("⚠️ Please upload and analyze your resume first.")

        return

    resume_text = st.session_state.get("resume_text", "")

    if not resume_text:

        st.error("Resume text not found.")

        return

    # ----------------------------------
    # Rewrite Button
    # ----------------------------------

    if st.button(
        "✨ Rewrite Resume",
        type="primary",
        use_container_width=True,
    ):

        with st.spinner("🤖 AI is rewriting your resume..."):

            try:

                integration = ResumeRewriterIntegration()

                rewritten = integration.rewrite_resume(resume_text)

                st.session_state["rewritten_resume"] = rewritten

            except Exception as e:

                st.error(str(e))

                return

    # ----------------------------------
    # Show Result
    # ----------------------------------

    if "rewritten_resume" in st.session_state:

        st.success("✅ Resume rewritten successfully!")

        rewritten_resume = st.session_state["rewritten_resume"]

        st.text_area(
            "📄 AI Rewritten Resume",
            rewritten_resume,
            height=600,
        )

        docx_file = create_resume_docx(rewritten_resume)

        st.download_button(
            label="⬇️ Download Resume (.docx)",
            data=docx_file,
            file_name="ATS_Rewritten_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )