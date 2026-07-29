import io
import os
import docx
from groq import Groq
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Create Groq client
client = Groq(
    api_key=os.getenv("GROQ_API_KEY")
)


def generate_full_rewritten_resume(
    resume_text: str,
    target_role: str = "Software Engineer / AI Specialist",
    tone: str = "High ATS & Impact Optimized"
) -> str:
    """
    Backend LLM engine to rewrite an ENTIRE candidate resume into a unified, high-ATS-optimized professional resume.
    Preserves all true candidate experience, dates, projects, and tech stack while upgrading bullet impact,
    XYZ achievement metrics, and ATS keywords.
    """
    if not resume_text.strip():
        return ""

    prompt = f"""
You are an elite executive resume writer and ATS optimization specialist.

Rewrite the following entire candidate resume into a single, cohesive, high-ATS-optimized professional resume.

Target Role Context: {target_role if target_role else "Tech Specialist"}
Optimization Tone: {tone}

Original Resume:
{resume_text}

Rules:
1. Format as a clean, complete, unified professional resume using standard sections:
   CANDIDATE SUMMARY
   TECHNICAL SKILLS
   PROFESSIONAL EXPERIENCE
   TECHNICAL PROJECTS
   EDUCATION & CREDENTIALS
2. Upgrade weak duty bullets into commanding achievement statements (XYZ formula: Action Verb + Task + Result).
3. Maximize ATS keyword alignment for {target_role}.
4. Preserve all true candidate experience, companies, dates, tools, and technical facts. Do NOT invent fake companies or fake experience.
5. Return ONLY the complete rewritten resume text cleanly formatted.
"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.3,
        max_tokens=1600
    )

    return response.choices[0].message.content.strip()


def create_docx_resume(full_resume_text: str) -> bytes:
    """
    Converts full rewritten resume text into a downloadable Microsoft Word (.docx) binary stream.
    """
    doc = docx.Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = docx.shared.Inches(0.75)
        section.bottom_margin = docx.shared.Inches(0.75)
        section.left_margin = docx.shared.Inches(0.75)
        section.right_margin = docx.shared.Inches(0.75)

    paragraphs = full_resume_text.split("\n")
    for line in paragraphs:
        line_clean = line.strip()
        if not line_clean:
            continue
            
        if line_clean.startswith("#") or (line_clean.isupper() and len(line_clean) < 35):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = docx.shared.Pt(10)
            p.paragraph_format.space_after = docx.shared.Pt(4)
            run = p.add_run(line_clean.lstrip("# ").strip())
            run.bold = True
            run.font.size = docx.shared.Pt(12)
            run.font.name = "Arial"
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = docx.shared.Pt(3)
            run = p.add_run(line_clean)
            run.font.size = docx.shared.Pt(10.5)
            run.font.name = "Arial"

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
